import datetime
import sys
import smtplib
import re
import smtplib
import os.path
import subprocess
import sys, getopt

#doc format
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

#email exchange
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

##################################### Guest Form as .docx

#command line arguments
inFile = ''
inOption = ''

try:
   opts, args = getopt.getopt(sys.argv[1:],"hi:f:",["ifile=","f="])
except getopt.GetoptError:
   print 'try.py -i <inputfile.txt> -f <["email" | "packet"]>'
   sys.exit(2)
for opt, arg in opts:
    if opt == '-h':
        print 'try.py -i <inputfile.txt> -f <["email" | "packet"]>'
        sys.exit()
    elif opt in ("-i", "--ifile"):
        inFile = arg
    elif opt in ("-f", "--f"):
         inOption = arg

print 'Input file is "', inFile
print 'Form: is "', inOption

subprocess.check_call("./parse_me.sh '%s'" %inFile,   shell=True)

#input file as an argument from the command line
outFile = re.sub('txt', 'out', inFile)

date = datetime.date.today()
villas = ['830\n','826\n','778\n','752\n']
phone_owners = ['14695697081','14087722615']
villa_names =['Rustica','Tortuga','Acqualina','Solana']
villa_name=""

#pull guest record
guest_info = []
with open(outFile, "r") as ins:
    	for line in ins:
        	guest_info.append(line)

#Guest form as .docx
document = Document()
document.add_picture('logo.png', width=Inches(2.0))

paragraph = document.add_paragraph()
paragraph.add_run("VISITORS/GUEST FORM").bold = True
style = document.styles['No Spacing']
font = style.font
font.underline = True
paragraph.add_run('					DATE: ').underline = False
paragraph.add_run(str(date))

pvilla_number = document.add_paragraph()
pvilla_number.add_run('VILLA # ').underline = False
if (str(guest_info[0]) ==  str(villas[0])):
	pvilla_number.add_run(str(villas[0])).underline = True
	villa_name=villa_names[0]
elif (str(guest_info[0]) == str(villas[1])):
	pvilla_number.add_run(str(villas[1])).underline = True
	villa_name=villa_names[1]
elif (str(guest_info[0]) == str(villas[2])):
	pvilla_number.add_run(str(villas[2])).underline = True
	villa_name=villa_names[2]
elif (str(guest_info[0]) == str(villas[3])):
	pvilla_number.add_run(str(villas[3])).underline = True
	villa_name=villa_names[3]

powner_name = document.add_paragraph()
powner_name.add_run('OWNER: ').underline = False
if (str(guest_info[0]) == str(villas[0])):
	powner_name.add_run('  Carlos Santa  ').underline = True
else:
	powner_name.add_run('  Serge  ').underline = True

pphone_owner = document.add_paragraph()
pphone_owner.add_run('PHONE OF OWNER: ').underline = False
if (str(guest_info[0]) == str(villas[0])):
	pphone_owner.add_run(str(phone_owners[0])).underline = True
else:
	pphone_owner.add_run(str(phone_owners[1])).underline = True

document.add_paragraph("")

prenter = document.add_paragraph()
prenter.add_run("VISITOR/RENTER").bold = True
#style = document.styles['No Spacing']
#font = style.font
font.underline = True

prenter_name = document.add_paragraph()
prenter_name.add_run('NAME: ').underline = False
guest_name = str(guest_info[1])
guest_name = re.sub('-', ' ', guest_name)
guest_name = re.sub('\n', '', guest_name)
prenter_name.add_run(guest_name).underline = True

pnumber_people = document.add_paragraph()
pnumber_people.add_run('# OF PEOPLE: ').underline = False
pnumber_people.add_run(guest_info[3]).underline = True
style = document.styles['No Spacing']
font = style.font
#font.underline = True
pnumber_people.add_run('ARRIVAL DATE: ').underline = False
pnumber_people.add_run(str(guest_info[4])).underline = True
pnumber_people.add_run('DEPARTURE DATE: ').underline = False
pnumber_people.add_run(str(guest_info[5])).underline = True

pphone_renter = document.add_paragraph()
pphone_renter.add_run('*TEL. PHONE NO. OF ONE OR MORE PEOPLE STAYING AT VILLA: ').underline = False
pphone_renter.add_run('\n1. ')
pphone_renter.add_run(guest_info[2]).underline = True
pphone_renter.add_run('\n*This is required')

pdisclaimer = document.add_paragraph()
pdisclaimer.add_run('Important: the attached form must be completed and given to guards or emailed to FWC Administrator at least 24 hours before anyone is staying at your villa when you are not there. fairwaycourts@outlook.com ').underline = False

#Save .docx using the guest's name and today's date
outputFile=str(guest_info[1]).rstrip("\n\r")+str(date)+"-guest.docx"
document.save(outputFile)

#parse language entry and cleanup the excess
language = str(guest_info[7])
#language = re.sub('US', '', language)
language = re.sub('-', ' ', language)
language = re.sub('\n', '', language)

sys.stdout.flush()

################################# Email exchange

welcome_packets = ['rustica-welcome-package-2018.pdf','rustica-paquete-bienvenida-2018.pages.pdf',
                   'tortuga-welcome-package-2018.pdf','tortuga-paquete-bienvenida-2018.pages.pdf',
                   'acqualina-welcome-package-2018.pdf','acqualina-welcome-package-2018.pdf',
                   'solana-welcome-package-2018.pdf','solana-paquete-bienvenida-2018.pages.pdf']

def sendemail(email, gmail_password, to, text):
	try:  
	    server = smtplib.SMTP('smtp.gmail.com', 587)#465
	    ##server.ehlo()
	    server.starttls()
	    server.login(email, gmail_password)
	    #text = msg.as_string()
	    server.sendmail(email, to, text)
	    ##server.close()
	    server.quit()
	    print "Success: Email sent to: "+str(to)
	except:  
	    print "Error: Email server error..."
	return

def composeemail(file_path, file_to_be_attached, email, to, subject, body, email_type ):
	try:
		msg = MIMEMultipart()
		msg['From'] = email
		msg['To'] = to
		msg['Subject'] = subject
		msg.attach(MIMEText(body, 'plain'))

		filename = os.path.basename(file_path+file_to_be_attached)
		attachment = open(filename, "rb")
		part = MIMEBase('application', 'octet-stream')
		part.set_payload(attachment.read())
		attachment.close()
		encoders.encode_base64(part)
		part.add_header('Content-Disposition', "attachment; filename= %s" % file_to_be_attached)

		if (str(email_type) == str("Guest Form")) :
			print("\n\n###################### GUEST FORM ###############################\n")
			#echo docx to console
			document = Document(outputFile)
			for para in document.paragraphs:
			    print(para.text)
			print("\n")
		else:
			print("\n\n###################### WELCOME PACKET ###############################\n")

		#echo at this point in order to review content without the attachment
		print str(msg)
		#now we can add the attachment to msg
		msg.attach(part)
		print"\nSuccess: Attachment: "+file_to_be_attached +" was found and attached!"		

		print (30 * '-')
		print ("   M A I N - M E N U")
		print (30 * '-')
		print ("1. Send "+email_type+" Email?:")
		print ("2. Exit")
		print (30 * '-')

		###########################
		## Robust error handling ##
		## only accept int       ##
		###########################
		## Wait for valid input in while...not ###
		is_valid=0
		 
		while not is_valid :
		       	try :
		                choice = int ( raw_input('Enter your choice [1-2] : ') )
		                is_valid = 1 ## set it to 1 to validate input and to terminate the while..not loop
		        except ValueError, e :
		                print ("'%s' is not a valid integer." % e.args[0].split(": ")[1])
		 
		### Take action as per selected menu-option ###
		if choice == 1:
		        print ("\nSending email...")
		        can_go_out = 1
		elif choice == 2:
		        print ("\nExit...")
		        can_go_out = 0
		        quit()
		else:
		        print ("\nInvalid option. Try again...")

		body = msg.as_string()

		if(can_go_out):
				sendemail(email, gmail_password, to, body)		 
	except:
		can_go_out = 0
		print"Attachment: "+file_to_be_attached +" was NOT found!"
		print"Error: Email was not sent..."
		quit()
	return

email = ""
gmail_password = ''

file_location = '/home/csanta/beach_houses/'

print(str(language))
print(str(guest_info[6]))

if (str(guest_info[0]) == str(villas[0])):#830
	if(str(language) == str("English US ") or (str(language) == str("English "))):
		villa_filename = welcome_packets[0]
	else:
		villa_filename = welcome_packets[1]
	villa_id = str(villas[0])
	villa_id = re.sub('\n', '', villa_id)
elif (str(guest_info[0]) == str(villas[1])):#826
	if(str(language) == str("English US ") or (str(language) == str("English "))):
		villa_filename = welcome_packets[2]
	else:
		print(str(language))
		villa_filename = welcome_packets[3]
	villa_id = str(villas[1])
	villa_id = re.sub('\n', '', villa_id)
elif (str(guest_info[0]) == str(villas[2])):#778
	if(str(language) == str("English US ") or (str(language) == str("English "))):
		villa_filename = welcome_packets[4]
	else:
		villa_filename = welcome_packets[5]
	villa_id = str(villas[2])
	villa_id = re.sub('\n', '', villa_id)
else:										#752
	if(str(language) == str("English ") or (str(language) == str("English "))):
		villa_filename = welcome_packets[6]
	else:
		villa_filename = welcome_packets[7]
	villa_id = str(villas[3])
	villa_id = re.sub('\n', '', villa_id)

if (str(inOption) == str("email")):
	#Guest Form Email
	to = ""
	subject = 'Guest Form: FWC #'+villa_id+': '+guest_name
	body = 'Hi Rose,\n\nPlease find in the attachment the guest authorization form for '+guest_name+'staying at villa FWC #'+villa_id+'. \n\nThanks,\nCarlos Santa'

	composeemail(file_location, outputFile, email, to, subject, body, "Guest Form")

elif (str(inOption) == str("packet")):
	#Welcome Packet Email
	to = guest_info[6]
	if ( (str(language) == "English ") or (str(language) == "English US ")):
		subject = 'Welcome packet to your villa in Palmas del Mar - FWC #'+villa_id
		body_wpacket = "Dear "+guest_name+":\n\nPlease find in the attachment your welcome packet to your rental. It has a lot of information about how to operate the villa including the access code as well as additional info on the surrounding areas in Palmas del Mar. The "+villa_name+" Beach House villa is located on the second floor within the Fairway Courts complex, unit #"+villa_id+".\n\nCheck in: "+guest_info[4]+"\nCheck out: "+guest_info[5]+"\nLet us know if you have any questions,\n\nCarlos Santa\n469-569-7081"
	else:
		subject = 'Paquete de bienvenida estadia Palmas del Mar - FWC #'+villa_id
		body_wpacket = "Saludos "+guest_name+":\n\nEncuentre en el atachado el documento de bienvenida en el complejo Fairway Courts en Palmas del Mar, el numero de la villa  es #"+villa_id+" y se encuentra en el segundo piso. Dentro del documento puede ver  el codigo de acceso a la villa para sacar las llaves de la cajita o lockbox.\n\nFecha de entrada: "+guest_info[4]+"\nFecha de salida: "+guest_info[5]+"\nAdemas del codigo para las llaves el documento contiene valiosa informacion de como operar la villa. Nos deja saber si tiene alguna pregunta.\n\nGracias,\n\nCarlos Santa\n469-569-7081"

	composeemail(file_location, villa_filename, email, to, subject, body_wpacket, "Welcome Packet")

else:
	print "Need to pass [email|packet] as a 2nd parameter"
	quit()

quit()
