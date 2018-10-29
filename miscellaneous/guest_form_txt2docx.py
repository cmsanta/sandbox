import datetime
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

date = datetime.date.today()
villas = ['830\n','826\n',
	  '778\n','752\n']
phone_owners = ['14695697081','14087722615']

array = []
with open("JessicaMartinez.txt", "r") as ins:
    for line in ins:
        array.append(line)

for idx, val in enumerate(array):
    print(str(val))

document = Document()
document.add_picture('logo.png', width=Inches(2.0))

paragraph = document.add_paragraph()
paragraph.add_run("VISITORS/GUEST FORM").bold = True
style = document.styles['No Spacing']
font = style.font
font.underline = True
paragraph.add_run('					DATE: ').underline = False
paragraph.add_run(str(date))

#p1 = document.add_paragraph()
#p1.add_run('test').font

pvilla_number = document.add_paragraph()
pvilla_number.add_run('VILLA # ').underline = False
pvilla_number.add_run(array[0]).underline = True
#document.add_paragraph("")

powner_name = document.add_paragraph()
powner_name.add_run('OWNER: ').underline = False
if (str(array[0]) == str(villas[0])):
	powner_name.add_run('  Carlos Santa  ').underline = True
else:
	powner_name.add_run('  Serge  ').underline = True
#document.add_paragraph("")

pphone_owner = document.add_paragraph()
pphone_owner.add_run('PHONE OF OWNER: ').underline = False
if (str(array[0]) == str(villas[0])):
	pphone_owner.add_run(str(phone_owners[0])).underline = True
else:
	pphone_owner.add_run(str(phone_onwers[1])).underline = True

document.add_paragraph("")

prenter = document.add_paragraph()
prenter.add_run("VISITOR/RENTER").bold = True
#style = document.styles['No Spacing']
#font = style.font
font.underline = True

prenter_name = document.add_paragraph()
prenter_name.add_run('NAME: ').underline = False
prenter_name.add_run(array[1]).underline = True
#document.add_paragraph("")

pnumber_people = document.add_paragraph()
pnumber_people.add_run('# OF PEOPLE: ').underline = False
pnumber_people.add_run(array[3]).underline = True
style = document.styles['No Spacing']
font = style.font
#font.underline = True
pnumber_people.add_run('ARRIVAL DATE: ').underline = False
pnumber_people.add_run(str(array[4])).underline = True
pnumber_people.add_run('DEPARTURE DATE: ').underline = False
pnumber_people.add_run(str(array[5])).underline = True

pphone_renter = document.add_paragraph()
pphone_renter.add_run('*TEL. PHONE NO. OF ONE OR MORE PEOPLE STAYING AT VILLA: ').underline = False
pphone_renter.add_run('		        1. ')
pphone_renter.add_run(array[2]).underline = True
pphone_renter.add_run('											   *This is required		        ')

pdisclaimer = document.add_paragraph()
pdisclaimer.add_run('Important: the attached form must be completed and given to guards or emailed to FWC Administrator at least 24 hours before anyone is staying at your villa when you are not there. fairwaycourts@outlook.com ').underline = False

document.save('guest.docx')

